import docx

handleName = docx.Document('D:\\projectsertif\\nama.docx')

for name in range(0,4):
    inp = handleName.paragraphs[name]
    nameSertif = inp.text
    doc = docx.Document('D:\\projectsertif\\editsertif.docx')
    p = doc.paragraphs[11]
    p.runs[0].text = f'{nameSertif}'
    doc.save(f'D:\\projectsertif\\hasil\\{nameSertif}.docx')